#!/usr/bin/env python3
"""Build the English KSC Global Scale-up Program application for DynaMOS."""

from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "1F4E79"
BLUE = "2E75B6"
PALE_BLUE = "D9EAF7"
VERY_PALE_BLUE = "EDF4FA"
LIGHT_GREY = "F2F2F2"
MID_GREY = "D9E1F2"
DARK_GREY = "666666"
YELLOW = "FFF2CC"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color="B4C6E7", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold=False, size=9.2, color=BLACK, align=None, highlight=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)


def set_run_font(run, *, size=9.5, bold=False, italic=False, color=BLACK, highlight=False) -> None:
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def style_paragraph(paragraph, *, before=0, after=3, line=1.08) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_footer_page_numbers(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0)
    run = p.add_run("Page ")
    set_run_font(run, size=8, color=DARK_GREY)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])

    run2 = p.add_run(" of ")
    set_run_font(run2, size=8, color=DARK_GREY)
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld_sep2 = OxmlElement("w:fldChar")
    fld_sep2.set(qn("w:fldCharType"), "separate")
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run2._r.extend([fld_begin2, instr2, fld_sep2, fld_end2])


def add_section_bar(doc: Document, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(1.1)
    table.columns[1].width = Cm(17.3)
    set_cell_shading(table.cell(0, 0), BLUE)
    set_cell_shading(table.cell(0, 1), NAVY)
    set_cell_text(table.cell(0, 0), number, bold=True, size=13, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), title, bold=True, size=12, color=WHITE)
    for cell in table.rows[0].cells:
        set_cell_border(cell, color=NAVY, size="2")
    p = doc.add_paragraph()
    style_paragraph(p, after=1)


def add_question(doc: Document, number: str, question: str, guidance: str, answer: str, *, review_fragments=()) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=2, after=1)
    r = p.add_run(f"{number}. {question}")
    set_run_font(r, size=10.2, bold=True, color=BLACK)

    p2 = doc.add_paragraph()
    style_paragraph(p2, after=2)
    r2 = p2.add_run(guidance)
    set_run_font(r2, size=8.4, italic=True, color=BLUE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, VERY_PALE_BLUE)
    set_cell_border(cell, color="B4C6E7", size="5")
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    p3 = cell.paragraphs[0]
    style_paragraph(p3, after=0, line=1.12)

    if not review_fragments:
        run = p3.add_run(answer)
        set_run_font(run, size=9.2)
    else:
        cursor = 0
        for fragment in review_fragments:
            idx = answer.find(fragment, cursor)
            if idx < 0:
                continue
            if idx > cursor:
                run = p3.add_run(answer[cursor:idx])
                set_run_font(run, size=9.2)
            run = p3.add_run(fragment)
            set_run_font(run, size=9.2, bold=True, highlight=True)
            cursor = idx + len(fragment)
        if cursor < len(answer):
            run = p3.add_run(answer[cursor:])
            set_run_font(run, size=9.2)

    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=1)


def add_review_note(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, YELLOW)
    set_cell_border(cell, color="D6B656", size="6")
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.05)
    r = p.add_run("APPLICANT REVIEW REQUIRED — ")
    set_run_font(r, size=8.4, bold=True, color="7F6000")
    r = p.add_run(
        "Yellow-highlighted fields were not available in the supplied materials. Confirm them before submission: legal English CEO name, funding details, quantitative traction, and the shareable pitch-deck URL. Remove this note after confirmation."
    )
    set_run_font(r, size=8.4, color="7F6000")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    add_footer_page_numbers(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    doc.core_properties.title = "KSC Global Scale-up Program 2026 — DynaMOS Application"
    doc.core_properties.subject = "Singapore market application"
    doc.core_properties.author = "RUNUP Co., Ltd."
    doc.core_properties.keywords = "DynaMOS, AI SCM, smart manufacturing, KSC, Singapore"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0)
    r = p.add_run("K-Startup Center (KSC) Global Scale-up Program 2026")
    set_run_font(r, size=16, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p2, after=8)
    r2 = p2.add_run("Application Form")
    set_run_font(r2, size=14, bold=True, color=BLUE)

    header = doc.add_table(rows=4, cols=4)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    widths = [Cm(3.1), Cm(6.1), Cm(3.1), Cm(6.1)]
    for row in header.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    labels = [
        (0, 0, "Company"), (0, 2, "Industry"),
        (1, 0, "Target market"), (1, 2, "CEO / Contact"),
        (2, 0, "Funding to date"), (2, 2, "Website"),
    ]
    for row, col, label in labels:
        set_cell_shading(header.cell(row, col), PALE_BLUE)
        set_cell_text(header.cell(row, col), label, bold=True, size=9.0, color=NAVY)

    set_cell_text(header.cell(0, 1), "RUNUP Co., Ltd. / DynaMOS", bold=True, size=9.2)
    set_cell_text(header.cell(0, 3), "AI-powered Smart Manufacturing / SCM-MES-ERP / Enterprise SaaS", size=8.8)
    set_cell_text(header.cell(1, 1), "Singapore", bold=True, size=9.4, color=BLUE)
    set_cell_text(
        header.cell(1, 3),
        "Hyun-tae Seo [CONFIRM LEGAL ENGLISH NAME] / support@runup.co.kr",
        size=8.6,
        highlight=True,
    )
    set_cell_text(
        header.cell(2, 1),
        "[CONFIRM: round, year, amount and investor(s)]",
        size=8.8,
        highlight=True,
    )
    set_cell_text(header.cell(2, 3), "https://runup.co.kr/en/", size=8.8, color=BLUE)

    merged = header.cell(3, 0).merge(header.cell(3, 3))
    set_cell_shading(merged, LIGHT_GREY)
    set_cell_text(
        merged,
        "Target market:   ☐ Shanghai       ☒ Singapore",
        bold=True,
        size=9.4,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_review_note(doc)

    add_section_bar(doc, "1", "Your company")

    add_question(
        doc,
        "1",
        "What does your company do?",
        "One or two plain sentences. What you make, and who it is for.",
        "DynaMOS is an AI-powered manufacturing operations and supply-chain platform for small and mid-sized manufacturers. It unifies master data, demand and production planning, purchasing, shop-floor execution, quality, shipping and closing, while a RAG/LLM copilot preserves operational know-how and gives context-aware answers to frontline and management teams.",
    )

    traction_placeholder = "[INSERT: 2025–2026 revenue, paying customers, active users and growth rate]"
    add_question(
        doc,
        "2",
        "How is it going so far?",
        "Revenue, customers or users, how fast you are growing. Rough numbers are fine — and it is okay if you are still early.",
        "RUNUP has built a working DynaMOS product covering the complete order-to-close workflow: master data/BOM, planning and MRP, purchasing, production, quality, shipping, closing and an embedded AI chatbot. The product screens and 2026–2027 roadmap are defined, and the team has implementation experience across manufacturing, logistics, energy and public-sector automation. Quantitative traction for the application: " + traction_placeholder + ".",
        review_fragments=(traction_placeholder,),
    )

    funding_placeholder = "[CONFIRM: bootstrapped or financing round, year, amount and investor(s)]"
    add_question(
        doc,
        "3",
        "How much have you raised, if any?",
        "Round, year, amount, investor. No problem if you have not raised — state that, and whether you are planning to.",
        funding_placeholder + ". We plan to use validated Singapore PoCs and a repeatable ASEAN go-to-market model as the basis for the next institutional financing round.",
        review_fragments=(funding_placeholder,),
    )

    add_question(
        doc,
        "4",
        "Why Singapore?",
        "What is the opportunity there — who would buy, and why is it a fit?",
        "Singapore is our preferred ASEAN beachhead because it offers an English-speaking enterprise environment, regional manufacturing headquarters and strong links among global manufacturers, government innovation agencies, research institutes and investors. DynaMOS fits advanced manufacturers and regional HQs that operate fragmented ERP, MES and APS workflows across ASEAN and need an explainable, knowledge-preserving AI operations layer. ATUM Ventures' advanced-manufacturing focus and network — including Enterprise Singapore, SGInnovate, A*STAR, NUS/NTUitive and corporate partners such as OMRON, Schneider Electric, Micron, ST Engineering, Panasonic and Murata — create a practical route to lighthouse PoCs and a local channel partner.",
    )

    doc.add_page_break()
    add_section_bar(doc, "2", "This market")
    intro = doc.add_paragraph()
    style_paragraph(intro, after=5)
    r = intro.add_run("If you have already started building something here, we would love to hear it. If you are starting fresh, that is fine too.")
    set_run_font(r, size=8.6, italic=True, color=BLUE)

    add_question(
        doc,
        "1",
        "What have you got going in this market already?",
        "Customers, revenue, partners, a local entity, PoCs, past visits, deals in the works. Names and dates help.",
        "No Singapore customer, local entity or confirmed local revenue is evidenced in the supplied materials, so we are treating the program as a focused market-entry sprint. We will enter with an English website, a working DynaMOS product and an implementation method that supports phased adoption. From July to September 2026, we will package a 6–8 week PoC around one of three measurable use cases: operational-knowledge retrieval, plan-versus-actual root-cause analysis, or quality-history analysis. We will prepare an English security/integration checklist and baseline metrics before the October on-site week, then convert qualified meetings into two PoC proposals.",
    )

    add_question(
        doc,
        "2",
        "Who do you already know there?",
        "Partners, investors, distributors, advisors — anyone you could lean on. Skip if none.",
        "No formal Singapore partner is confirmed in the supplied materials. Through ATUM and KSC, we will build working relationships with Enterprise Singapore/SGInnovate, A*STAR, NUS/NTUitive, advanced-manufacturing companies in ATUM's network, local SI/ERP/cloud partners, and industrial-AI investors. Our immediate priority is one technical integration partner and two corporate PoC sponsors.",
    )

    add_question(
        doc,
        "3",
        "What makes you stand out?",
        "Anything worth knowing — awards, patents, press, partnerships, certifications.",
        "Most SME manufacturers keep ERP/MES/APS data separate from manuals, logs, meeting records, change histories and people-dependent know-how. DynaMOS puts both structured operations and unstructured knowledge into one operating layer. Its RAG/LLM assistant retrieves past records, interprets them by user role, process and current state, and explains why an outcome occurred and what action to take; the same platform executes the seven-stage manufacturing workflow. This preserves knowledge through staff changes, shortens handovers and enables phased adoption rather than a high-risk rip-and-replace. The roadmap extends to Hybrid RAG, workflow/job scheduling, agentic AI with MCP/A2A, user-defined tuning, and knowledge-graph alerts.",
    )

    doc.add_page_break()
    add_section_bar(doc, "3", "The program")

    p = doc.add_paragraph()
    style_paragraph(p, before=2, after=1)
    r = p.add_run("1. What do you want to get out of this?")
    set_run_font(r, size=10.2, bold=True)
    p2 = doc.add_paragraph()
    style_paragraph(p2, after=3)
    r2 = p2.add_run("Your top 3 goals for the program (Jul–Dec 2026). The more concrete, the better.")
    set_run_font(r2, size=8.4, italic=True, color=BLUE)

    goals = doc.add_table(rows=4, cols=3)
    goals.alignment = WD_TABLE_ALIGNMENT.CENTER
    goals.autofit = False
    goals.columns[0].width = Cm(1.0)
    goals.columns[1].width = Cm(6.3)
    goals.columns[2].width = Cm(11.0)
    for idx, text in enumerate(("#", "Goal", "What success looks like")):
        set_cell_shading(goals.cell(0, idx), PALE_BLUE)
        set_cell_text(goals.cell(0, idx), text, bold=True, size=9.1, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else None)

    goal_rows = [
        (
            "1",
            "Secure Singapore lighthouse PoCs",
            "By 15 Dec 2026: 15 qualified business meetings, 5 technical discovery sessions, and 2 signed PoC agreements or LOIs, including at least one advanced manufacturer or ASEAN regional HQ.",
        ),
        (
            "2",
            "Localize and prove repeatable deployment",
            "English product/IR/security pack by 31 Aug; a 6–8 week PoC package by 30 Sep; one live data integration and one AI use case deployed, targeting at least a 30% reduction in time spent retrieving operational information or completing handovers.",
        ),
        (
            "3",
            "Build an ASEAN channel and financing pipeline",
            "By Dec 2026: one MoU with a Singapore SI/ERP/cloud partner, 8 investor meetings, 3 follow-up diligence requests, and a qualified 12-month Singapore/ASEAN sales pipeline.",
        ),
    ]
    for row_idx, (num, goal, success) in enumerate(goal_rows, start=1):
        fill = VERY_PALE_BLUE if row_idx % 2 else WHITE
        for col in range(3):
            set_cell_shading(goals.cell(row_idx, col), fill)
        set_cell_text(goals.cell(row_idx, 0), num, bold=True, size=9.1, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(goals.cell(row_idx, 1), goal, bold=True, size=8.8)
        set_cell_text(goals.cell(row_idx, 2), success, size=8.5)
        goals.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_question(
        doc,
        "2",
        "What do you need from us?",
        "For example: introductions to investors, a PoC with a local company, entity setup or a distribution partner.",
        "1) Warm introductions to 10–15 advanced manufacturers or regional HQs and two potential PoC sponsors, especially within ATUM's OMRON, Schneider Electric, Micron, ST Engineering, Panasonic and Murata network. 2) A Singapore technical/commercial partner for integration, cybersecurity, cloud/data-residency and enterprise procurement. 3) Investor introductions in industrial AI and enterprise SaaS, English IR rehearsal, and local pricing/contracting guidance. 4) Follow-up support to convert PoCs into paid rollouts and advice on the timing and structure of a Singapore entity.",
    )

    p_links = doc.add_paragraph()
    style_paragraph(p_links, before=3, after=2)
    r = p_links.add_run("Links (optional)")
    set_run_font(r, size=10.2, bold=True)

    links = doc.add_table(rows=2, cols=2)
    links.alignment = WD_TABLE_ALIGNMENT.CENTER
    links.autofit = False
    links.columns[0].width = Cm(4.4)
    links.columns[1].width = Cm(13.9)
    set_cell_shading(links.cell(0, 0), PALE_BLUE)
    set_cell_shading(links.cell(1, 0), PALE_BLUE)
    set_cell_text(links.cell(0, 0), "Pitch deck", bold=True, size=9.0, color=NAVY)
    set_cell_text(links.cell(1, 0), "Demo / press", bold=True, size=9.0, color=NAVY)
    set_cell_text(links.cell(0, 1), "[INSERT SHAREABLE ENGLISH PITCH-DECK URL]", size=8.8, highlight=True)
    set_cell_text(links.cell(1, 1), "https://runup.co.kr/en/solution/dynamos/", size=8.8, color=BLUE)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    note = doc.add_paragraph()
    style_paragraph(note, before=4, after=0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Submission check: replace all yellow-highlighted text and remove the applicant review note before emailing the application.")
    set_run_font(r, size=8.2, italic=True, color=DARK_GREY)

    return doc


def main() -> None:
    output_dir = Path("build")
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / "DynaMOS_KSC_2026_Singapore_Application_Draft.docx"
    doc = build_document()
    doc.save(output)

    with zipfile.ZipFile(output, "r") as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
        missing = required - names
        if missing:
            raise RuntimeError(f"Generated DOCX is missing required parts: {sorted(missing)}")
        bad = archive.testzip()
        if bad is not None:
            raise RuntimeError(f"Generated DOCX has a corrupt ZIP member: {bad}")

    print(output)


if __name__ == "__main__":
    main()
